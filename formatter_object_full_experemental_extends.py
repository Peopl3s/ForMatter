from docx import Document
from docx.shared import Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE
from docx.shared import RGBColor
from docx.enum.text import WD_LINE_SPACING
import PySimpleGUI as sg
from itertools import count, takewhile
import json
import copy
import re

class DocFormatException(Exception):
    def __init__(self, text):
        self.txt = text
        
class TemplateFormatException(Exception):
    def __init__(self, text):
        self.txt = text

class Formatter():
    def __init__(self, windowGui = None):
        self.docAligment ={'По левому краю' : WD_ALIGN_PARAGRAPH.LEFT, 
			'По центру' : WD_ALIGN_PARAGRAPH.CENTER,
			'По правому краю' : WD_ALIGN_PARAGRAPH.RIGHT, 
			'По ширине' : WD_ALIGN_PARAGRAPH.JUSTIFY}
        
        self.docSpaceLine = {'1.5': WD_LINE_SPACING.ONE_POINT_FIVE,
			    'Двойной':WD_LINE_SPACING.DOUBLE,
			    'Одинарный':WD_LINE_SPACING.SINGLE,
			    'Минимум':WD_LINE_SPACING.AT_LEAST}
    
        self.keys = ['-aligment-', "-linespace-", "-redline-", "-before-",
                "-after-", "-intright-", "-intleft-", '-font-', '-sections-']
        self.font = ['-fontname-', '-fontsize-', '-bold-', '-italic-', '-underline-', '-fontcolor-']
        self.sections = ["-top-", "-bottom-", "-left-", "-right-"]
        self.__document = None
        self.__values = None
        self.__window = windowGui

    def updateValues(self, newValues):
        self.__values = newValues

    def __isHexColor(self, colorString):
        return re.match(r'#[\da-f]{6}', colorString, re.IGNORECASE) != None
        
    def __paragraphAlign(self, paragraph_format, alignParametr):
        '''Функция устанвливает выравнивание параграфа
           Arguments:
           paragraph_format -- Provides access to paragraph formatting such (docx.text.parfmt)
           alignParametr -- объект перечисления WD_PARAGRAPH_ALIGNMENT, устанавливающий выравнивание 
        '''
        paragraph_format.alignment = alignParametr
        
    def __lineSpacing(self, paragraph_format, spacingRule):
        paragraph_format.line_spacing_rule = spacingRule
        
    def __spaceBeforeAfter(self, paragraph_format, *, after_metrics, before_metrics,
                     right_metrics, left_metrics):
        paragraph_format.space_after = after_metrics
        paragraph_format.space_before = before_metrics
        paragraph_format.left_indent = left_metrics
        paragraph_format.right_indent = right_metrics

    def __firstLineIndent(self, paragraph_format, metrics):
        paragraph_format.first_line_indent = metrics

    def __docSections(self, sections, *, top, bottom, left, right):
        for section in sections:
            section.top_margin = Cm(top)
            section.bottom_margin = Cm(bottom)
            section.left_margin = Cm(left)
            section.right_margin = Cm(right)

    def __changeFont(self, runsCollection, *, fontName, fontSize, fontType, fontColor):
        if self.__isHexColor(fontColor):
            for run in runsCollection:
                font = run.font
                font.name = fontName
                font.bold, font.italic, font.underline = fontType
                font.color.rgb = RGBColor.from_string(fontColor[1:])
                font.size = Pt(fontSize)
        else:
            raise ValueError()
              
    def __setUp(self, tables, paragraphs, sections, aligment, spaceLineRule, fontName_,
              fontSize_, fontColor_, fontStyles_, redLine,
              top_, bottom_, left_, right_,
              afterSpace, beforeSpace, rightSpace, leftSpace):

        for paragraph in paragraphs:
            paragraph_format = paragraph.paragraph_format
            try:
                self.__changeFont(paragraph.runs, fontName = fontName_, fontSize = fontSize_,
                           fontType = fontStyles_, fontColor = fontColor_)
                
                self.__paragraphAlign(paragraph_format, aligment)

                self.__lineSpacing(paragraph_format, spaceLineRule)
            
                self.__firstLineIndent(paragraph_format, Cm(float(redLine)))

                self.__docSections(sections, top = float(top_), 
                                            bottom = float(bottom_), 
                                            left = float(left_), 
                                            right = float(right_))
            
                self.__spaceBeforeAfter(paragraph_format,
                                        after_metrics = Pt(float(afterSpace)),
                                        before_metrics = Pt(float(beforeSpace)),
                                        right_metrics = Cm(float(rightSpace)), 
                                        left_metrics = Cm(float(leftSpace)))
            except ValueError as verr:
                print(verr)
                sg.Popup("Ошибка! В полях введены некорректные данные")
                return 
            
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_format = paragraph.paragraph_format
                        self.__changeFont(paragraph.runs, fontName = fontName_, 
							fontSize = fontSize_,
							fontType = fontStyles_, 
							fontColor = fontColor_)
                
                        self.__paragraphAlign(paragraph_format, aligment)

                        self.__lineSpacing(paragraph_format, spaceLineRule)

                        self.__firstLineIndent(paragraph_format, Cm(float(redLine)))

                        self.__docSections(sections, top = float(top_), 
							bottom = float(bottom_), 
							left = float(left_), 
							right = float(right_))
                    
                        self.__spaceBeforeAfter(paragraph_format, 
								after_metrics = Pt(float(afterSpace)),
								before_metrics = Pt(float(beforeSpace)),
								right_metrics = Cm(float(rightSpace)), 
								left_metrics = Cm(float(leftSpace)))
    def __formSaveFilePath(self, pathFile, *, postfix = '_'):
        lastDot = pathFile.rfind('.')
        savePath = pathFile[:lastDot] + postfix + pathFile[lastDot:]
        return savePath
        
    def goFormat(self):
        pathFile = self.__values['-docfile-']
        if pathFile.endswith('.docx') or pathFile.endswith('.doc'):
            self.__document = Document(str(self.__values['-docfile-']))
        else:
            raise DocFormatException("Docx/doc format error")
        try:
            aligment = self.docAligment[self.__values['-aligment-']]
            spaceLineRule = self.docSpaceLine[self.__values['-linespace-']]
            fontName = self.__values['-fontname-']
            fontSize = self.__values['-fontsize-']
            fontColor = self.__values['-fontcolor-']
            fontStyles = [self.__values['-bold-'], self.__values['-italic-'],
                          self.__values['-underline-']]
            redLine = self.__values['-redline-']
            topSection = self.__values['-top-']
            bottomSection = self.__values['-bottom-']
            leftSection = self.__values['-left-']
            rightSection = self.__values['-right-']
            afterSpace = self.__values['-after-']
            beforeSpace = self.__values['-before-']
            rightSpace = self.__values['-intright-']
            leftSpace = self.__values['-intleft-']
        except Exception as e:
            print(e)
            sg.Popup("Ошибка! В полях введены некорректные данные")
            return 

        self.__setUp(self.__document.tables, self.__document.paragraphs, self.__document.sections, 
					aligment, spaceLineRule, fontName, fontSize, fontColor, 
					fontStyles, redLine, topSection, bottomSection, leftSection, 
					rightSection, afterSpace, beforeSpace, rightSpace, leftSpace)
        self.__document.save(self.__formSaveFilePath(pathFile, postfix = '_'))
                             
    def __parseJson(self, jsonFile):
        jsObj = None
        try:
            jsObj = open(jsonFile, "r",  encoding='utf-8')
            pObj = json.load(jsObj)
        except ValueError as VErr:
            print(VErr)
            raise TemplateFormatException("Error not a json")
        except IOError as IOErr:
            print(IOErr)
            raise TemplateFormatException("Error not a json")
        finally:
            if jsObj:
                jsObj.close()
        return  pObj

    def saveTemplate(self):
        tamplate = dict()
        for key in self.keys:
            if key in ('-font-', '-sections-'):
                tamplate[key] = dict()
                newList = self.font if key == '-font-' else self.sections
                for nested_kay in newList:
                    tamplate[key][nested_kay]=self.__values[nested_kay]
            else:
                tamplate[key]=self.__values[key]
            

        saveFilePath = self.__values['-saved-']
        try:
            with open(saveFilePath, 'w',encoding='utf-8') as jsonFile:
                jsonFile.write(json.dumps(tamplate))
        except FileNotFoundError as e:
            print(e)
            sg.Popup("Ошибка! Файл не сохранён!")
            return 
        
    def uploadTemplate(self):
        try:
            settings = self.__parseJson(self.__values['-template-'])
        except TemplateFormatException as et:
            print(et)
            sg.Popup("Ошибка! Можно использовать только json документы")
            return
        
        for key, value in settings.items():
            if key not in ('-font-', '-sections-', '-templatename-'):
                self.__window[key].update(settings[key])
            elif key in ('-font-', '-sections-'):
                    for nasted_key, nasted_values in settings[key].items():
                        self.__window[nasted_key].update(settings[key][nasted_key])

def frange(start, stop, step):
    return takewhile(lambda x: x < stop, count(start, step))


def createLayout():
    layout = [
		[sg.Text('Исходный документ'), sg.InputText(key='-docfile-', size=(60, 1), pad=((3,0),0)), sg.FileBrowse("Выбрать..", key ='-choicefile-', pad=((5,0),0))],
                [sg.Frame(layout=[
                 [sg.Text('Тип шрифта:'), sg.InputCombo(('Times New Roman','Calibri','Tahoma','Franklin Gothic Medium','Comic Sans','Segoe','Align'), size=(23, 1), key='-fontname-', pad=((2,0),0)),
                  sg.Text('Размер:', pad=((10,0),0)), sg.Spin([i for i in range(1,72)], initial_value=12, key ='-fontsize-', size=(4, 1)),
                  sg.ColorChooserButton("Выберите цвет", target = '-fontcolor-',key='-color-', size=(15, 1), pad=((18,0),0)), 
                  sg.InputText(enable_events=True, key='-fontcolor-', size = (12, 1), pad=((4,0),0), justification='center')],
                 [sg.Text('Выравнивание:'),sg.InputCombo(('По левому краю', 'По центру', 'По правому краю','По ширине'), size=(20, 1), key='-aligment-', pad=((6,0),0)),
                  sg.Checkbox('Жирный', key='-bold-', font=('Segoe', 10, 'bold'), pad=((10,0),0)), 
                  sg.Checkbox('Курсив', key='-italic-', font=('Segoe', 10, 'italic'), pad=((37,0),0)),
                  sg.Checkbox('Подчёркнутый', key = '-underline-', font=('Segoe', 10, 'underline'), pad=((37,0),0))]], title='Текст', font=('Segoe', 12, 'bold'))],
                [sg.Frame(layout=[
                 [sg.Text('Междустрочный интервал:'), sg.InputCombo(('1.5', 'Двойной', 'Одинарный','Минимум'), size=(11, 1), key='-linespace-'),
                  sg.Text('Интервал перед:', size=(13, 1), pad=((10,0),0)), sg.Spin([str(i) for i in frange(0,72,0.25)], initial_value=0, key ='-before-', size = (5,1)),
                  sg.Text('Интервал после:', size=(14, 1), pad=((5,0),0)), sg.Spin([str(i) for i in frange(0,72,0.25)], initial_value=0, key ='-after-', size = (5,1))],
                 [sg.Text('Начало абзаца:'), sg.Spin([str(i) for i in frange(0,72,0.25)], initial_value=1.25,key ='-redline-', size=(5, 1), pad=((5,0),0)),
                  sg.Text('Интервал cлева:', size=(13, 1), pad=((120,0),0)), sg.Spin([str(i) for i in frange(0,72,0.25)], initial_value=0, key ='-intleft-', size = (5,1)),
                  sg.Text('Интервал cправа:', size=(14, 1), pad=((5,0),0)), sg.Spin([str(i) for i in frange(0,72,0.25)], initial_value=0, key ='-intright-', size = (5,1))]], title='Абзац', font=('Segoe', 12, 'bold'))],
                [sg.Frame(layout=[
                [sg.Text('Верхнее:', size=(7, 1)), sg.Spin([str(i) for i in frange(0,72,0.25)], initial_value=0, key ='-top-', size=(5, 1)),
                 sg.Text('Нижнее:', size=(6, 1), pad=((5,2),0)), sg.Spin([str(i) for i in frange(0,72,0.25)], initial_value=0, key ='-bottom-', size=(5, 1))],
                [sg.Text('Левое:', size=(7, 1)), sg.Spin([str(i) for i in frange(0,72,0.25)], initial_value=0, key ='-left-', size=(5, 1)),
                 sg.Text('Правое:', size=(6, 1), pad=((5,2),9)), sg.Spin([str(i) for i in frange(0,72,0.25)], initial_value=0, key ='-right-', size=(5, 1))]], title='Поля', font=('Segoe', 12, 'bold')),
                 sg.Frame(layout=[
                 [sg.FileBrowse("Загрузить шаблон", key='-template-', size=(15, 0)), 
                  sg.InputText(enable_events=True, key='-example-', size=(31, 0), pad=((3,0),0))],
                 [sg.FileSaveAs("Сохранить шаблон", file_types=(('ALL Files', '*.json'),), target='-saved-',key='-save-', size=(15, 0), pad=((1,1),5)),
                  sg.InputText(enable_events=True, key = '-saved-', size=(31, 0), pad=((3,0),0))]], title='Шаблоны', font=('Segoe', 12, 'bold'), pad=((5,0),0))],
                [sg.Button("Применить", size=(14, 2), pad=((0,5),4)),
                 sg.Button("Закрыть", size=(14, 2), pad=((5,0),4))]
             ]
    return layout

def main():
    sg.LOOK_AND_FEEL_TABLE['MyTheme'] = {'BACKGROUND': '#1a1a1a',
                                            'TEXT': '#ffffff',
                                            'INPUT': '#000000',
                                            'TEXT_INPUT': '#ffffff',
                                            'SCROLL': '#ff1111',
                                            'BUTTON': ('#ffffff', '#303030'),
                                            'PROGRESS': ('#01826B', '#D0D0D0'),
                                            'BORDER': 0, 'SLIDER_DEPTH': 0, 'PROGRESS_DEPTH': 0,
                                            }
                                            
    sg.theme('MyTheme')
    sg.SetOptions(element_padding = (1,4))    
    layout = createLayout()
    window = sg.Window('Formatter', layout)
    formatter = Formatter(window)                   
    while True:
        event, values = window.read()
        formatter.updateValues(values)
        if event in (None, 'Закрыть'):
            break
        if event == '-saved-':
            formatter.saveTemplate()
        if event == '-example-':
            formatter.uploadTemplate()
        if event == 'Применить':
            try:
                formatter.goFormat()
            except DocFormatException as ed:
                print(ed)
                sg.Popup("Ошибка! Можно использовать только docx/doc документы")
    window.close()
    
if __name__ == '__main__':
    try:
        main()
    except KeyboardInterrupt:
        exit()
